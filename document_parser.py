import os
import re
import html2text

# ── Encoding detection helper ─────────────────────────────────────────────

def read_text_file(file_path: str) -> str:
    """
    Try multiple encodings — never crash on encoding mismatch.
    Latin-1 accepts any byte sequence — guaranteed last resort.
    """
    for encoding in ['utf-8', 'utf-8-sig', 'cp1252', 'latin-1']:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    # Absolute fallback
    with open(file_path, 'rb') as f:
        return f.read().decode('utf-8', errors='ignore')
    
# ── DOCX parser ───────────────────────────────────────────────────────────

def parse_docx(file_path: str, filename: str) -> dict:
    """
    Extract text from Word documents preserving heading hierarchy.
    Tables extracted as readable formatted text.
    Bold text preserved with ** markers as emphasis signals.
    """
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    doc = Document(file_path)
    sections = []
    current_section = []
    current_heading = ""

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        text = para.text.strip()

        if not text:
            continue

        # Preserve bold runs with ** markers
        formatted_parts = []
        for run in para.runs:
            if run.bold and run.text.strip():
                formatted_parts.append(f"**{run.text}**")
            else:
                formatted_parts.append(run.text)
        formatted_text = ''.join(formatted_parts).strip() or text

        # Heading detection — split sections at heading boundaries
        is_heading = 'Heading' in style_name or 'Title' in style_name

        if is_heading:
            # Save previous section
            if current_section:
                section_text = '\n'.join(current_section)
                if current_heading:
                    sections.append(f"{current_heading}\n{section_text}")
                else:
                    sections.append(section_text)
                current_section = []
            current_heading = formatted_text

        else:
            # Body text — add to current section
            if 'List' in style_name:
                current_section.append(f"• {formatted_text}")
            else:
                current_section.append(formatted_text)

    # Don't forget last section
    if current_section:
        section_text = '\n'.join(current_section)
        if current_heading:
            sections.append(f"{current_heading}\n{section_text}")
        else:
            sections.append(section_text)

    # Extract tables as formatted text
    table_texts = []
    for i, table in enumerate(doc.tables):
        table_lines = [f"[Table {i+1}]"]
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # skip empty rows
                table_lines.append(' | '.join(cells))
        if len(table_lines) > 1:
            table_texts.append('\n'.join(table_lines))

    # Combine sections and tables
    all_content = '\n\n'.join(sections)
    if table_texts:
        all_content += '\n\n' + '\n\n'.join(table_texts)

    # Label first line for name/title queries
    first_meaningful = next(
        (p.text.strip() for p in doc.paragraphs if p.text.strip()), ""
    )
    if first_meaningful:
        all_content = f"Document title or name: {first_meaningful}\n\n" + all_content

    return {
        "text": all_content,
        "filename": filename,
        "type": "docx",
        "sections_found": len(sections),
        "tables_found": len(table_texts)
    }


# ── TXT parser ────────────────────────────────────────────────────────────

def parse_txt(file_path: str, filename: str) -> dict:
    """
    Plain text extraction with encoding detection and cleanup.
    """
    text = read_text_file(file_path)

    # Normalise whitespace — collapse multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = text.strip()

    if not text:
        return {"error": "File appears to be empty"}

    # Label first non-empty line
    lines = text.split('\n')
    first_line = next((l.strip() for l in lines if l.strip()), "")
    if first_line:
        text = f"Document title: {first_line}\n\n" + text

    return {
        "text": text,
        "filename": filename,
        "type": "txt"
    }


# ── Markdown parser ───────────────────────────────────────────────────────

def parse_markdown(file_path: str, filename: str) -> dict:
    """
    Parse Markdown preserving heading structure and code blocks.
    Headings become section boundaries.
    Code blocks extracted as atomic units — never split.
    """
    text = read_text_file(file_path)

    # Extract and remove YAML front matter
    front_matter = ""
    front_matter_match = re.match(r'^---\n(.*?)\n---\n', text, re.DOTALL)
    if front_matter_match:
        front_matter_raw = front_matter_match.group(1)
        text = text[front_matter_match.end():]

        # Parse key fields from front matter
        title_match = re.search(r'title:\s*["\']?(.+?)["\']?\s*$',
                                front_matter_raw, re.MULTILINE)
        author_match = re.search(r'author:\s*["\']?(.+?)["\']?\s*$',
                                 front_matter_raw, re.MULTILINE)

        meta_parts = []
        if title_match:
            meta_parts.append(f"Document title: {title_match.group(1).strip()}")
        if author_match:
            meta_parts.append(f"Author: {author_match.group(1).strip()}")
        if meta_parts:
            front_matter = '\n'.join(meta_parts) + '\n\n'

    # Protect code blocks — mark them so chunker never splits them
    code_blocks = {}
    def protect_code_block(match):
        key = f"__CODE_BLOCK_{len(code_blocks)}__"
        lang = match.group(1) or "text"
        code_blocks[key] = f"[Code block — language: {lang}]\n{match.group(2)}"
        return f"\n{key}\n"

    text = re.sub(
        r'```(\w*)\n(.*?)```',
        protect_code_block,
        text,
        flags=re.DOTALL
    )

    # Convert heading markers to section labels
    # # Heading → [SECTION: Heading]
    text = re.sub(r'^#{1,2}\s+(.+)$', r'\n[SECTION: \1]\n', text, flags=re.MULTILINE)
    text = re.sub(r'^#{3,6}\s+(.+)$', r'\n[\1]\n', text, flags=re.MULTILINE)

    # Restore code blocks
    for key, code in code_blocks.items():
        text = text.replace(key, code)

    # Clean up
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Prepend front matter
    final_text = front_matter + text.strip()

    return {
        "text": final_text,
        "filename": filename,
        "type": "markdown",
        "code_blocks_found": len(code_blocks)
    }


# ── EPUB parser ───────────────────────────────────────────────────────────

def parse_epub(file_path: str, filename: str) -> dict:
    """
    Extract EPUB chapters in spine order.
    Each chapter title becomes a section header.
    HTML converted to clean text via html2text.
    """
    try:
        import ebooklib
        from ebooklib import epub
        import html2text
    except ImportError:
        return {"error": "ebooklib/html2text not installed. Run: pip install ebooklib html2text"}

    book = epub.read_epub(file_path)

    # Configure html2text converter
    converter = html2text.HTML2Text()
    converter.ignore_links = True       # don't clutter text with URLs
    converter.ignore_images = True      # no alt text from images
    converter.body_width = 0            # no line wrapping
    converter.ignore_emphasis = False   # keep bold/italic signals

    # Extract metadata
    meta_parts = []
    title = book.get_metadata('DC', 'title')
    author = book.get_metadata('DC', 'creator')
    if title:
        meta_parts.append(f"Book title: {title[0][0]}")
    if author:
        meta_parts.append(f"Author: {author[0][0]}")
    metadata_text = '\n'.join(meta_parts) + '\n\n' if meta_parts else ""

    # Extract chapters in spine order
    chapters = []
    chapter_num = 0

    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            html_content = item.get_content().decode('utf-8', errors='ignore')

            # Convert HTML to clean text
            chapter_text = converter.handle(html_content).strip()

            if not chapter_text or len(chapter_text) < 50:
                continue  # skip empty chapters (TOC pages, etc.)

            chapter_num += 1

            # Try to extract chapter title from first heading
            title_match = re.search(r'^#+\s+(.+)$', chapter_text, re.MULTILINE)
            if title_match:
                chapter_title = title_match.group(1).strip()
            else:
                chapter_title = f"Chapter {chapter_num}"

            chapters.append(f"[{chapter_title}]\n{chapter_text}")

    if not chapters:
        return {"error": "No readable chapters found in EPUB"}

    full_text = metadata_text + '\n\n'.join(chapters)

    return {
        "text": full_text,
        "filename": filename,
        "type": "epub",
        "chapters_found": len(chapters)
    }